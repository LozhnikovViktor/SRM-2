# tenders/views.py
import json
from django.shortcuts import render, redirect, get_object_or_404
from django.template.loader import render_to_string
from django.db.models import F, Sum, Count, Avg, Q, ExpressionWrapper, DecimalField
from django.db import models as db_models  # ← Добавлено для models.Q
from django.views.generic import ListView, CreateView, DeleteView, UpdateView, DetailView, View
from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib.auth.models import User  # ← Добавлено для User.objects
from django.urls import reverse_lazy
from django.contrib.auth import login
from django.contrib.auth.forms import UserCreationForm
from django.contrib import messages
from django.utils import timezone
from django.http import HttpResponse
from datetime import timedelta
from collections import defaultdict
from .forms import CustomUserCreationForm, SearchTenderForm, ClientForm, TenderForm, TenderDocumentForm
from .models import AuditLog
from django.http import JsonResponse
from .forms import CustomUserCreationForm, SearchTenderForm, ClientForm, TenderForm, TenderDocumentForm, TenderFilterForm
from django.views.decorators.http import require_POST
from .forms import CommentForm
from django.contrib.auth.decorators import login_required
from django.db.models import Sum, F, Count, Q
from .models import Tender, Client, TenderDocument, AuditLog, Comment, TenderStatus
from .tenderplan_api import api_client
from .forms import TenderplanSearchForm
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from datetime import datetime, timedelta
from django.db.models import Sum, Count, F, Q
from django.contrib.auth.decorators import user_passes_test
from rest_framework import viewsets
from .serializers import TenderSerializer, ClientSerializer
from .models import ChatRoom, ChatMessage, Tender, Client
from .models import Product, ShipmentRequest, ShipmentRequestItem
from .forms import ShipmentRequestForm, ShipmentItemFormSet
from django.core.mail import send_mail
from django.conf import settings
from django.http import JsonResponse








# Импорты из приложения

from .utils import export_tenders_to_excel, export_dashboard_stats_to_excel, search_tenders_on_zakupki
from .models import Tender, Client, TenderDocument

def get_client_ip(request):
    """Получает IP-адрес клиента"""
    x_forwarded_for = request.META.get('HTTP_X_FORWARDED_FOR')
    if x_forwarded_for:
        ip = x_forwarded_for.split(',')[0]
    else:
        ip = request.META.get('REMOTE_ADDR')
    return ip
# ============================================
# 🔹 АВТОРИЗАЦИЯ
# ============================================
class RegisterView(CreateView):
    form_class = CustomUserCreationForm
    template_name = 'tenders/register.html'
    success_url = reverse_lazy('tenders:list')

    def form_valid(self, form):
        user = form.save()
        login(self.request, user)
        return super().form_valid(form)


# ============================================
# 🔹 ТЕНДЕРЫ
# ============================================

from .forms import TenderFilterForm


class TenderListView(LoginRequiredMixin, ListView):
    """Расширенный список тендеров с фильтрацией"""
    model = Tender
    template_name = 'tenders/tender_list.html'
    context_object_name = 'tenders'
    paginate_by = 10

    def get_form(self):
        """Создаёт форму с данными из GET-параметров"""
        return TenderFilterForm(self.request.GET)

    def get_queryset(self):
        queryset = Tender.objects.select_related('client', 'author', 'client__manager')
        form = self.get_form()
        
        if not form.is_valid():
            return queryset
        
        data = form.cleaned_data
        
        # 🔹 Текстовый поиск (по нескольким полям сразу)
        if data.get('search'):
            search = data['search']
            queryset = queryset.filter(
                Q(customer_name__icontains=search) |
                Q(executor_name__icontains=search) |
                Q(winner__icontains=search) |
                Q(comment__icontains=search) |
                Q(client__name__icontains=search) |
                Q(client__inn__icontains=search)
            )
        
        #  Фильтр по статусу (мульти-выбор)
        if data.get('status'):
            queryset = queryset.filter(status__in=data['status'])
        
        # 🔹 Диапазон сумм
        if data.get('amount_from'):
            queryset = queryset.filter(initial_amount__gte=data['amount_from'])
        if data.get('amount_to'):
            queryset = queryset.filter(initial_amount__lte=data['amount_to'])
        
        # 🔹 Диапазон дат дедлайна
        if data.get('deadline_from'):
            queryset = queryset.filter(deadline__date__gte=data['deadline_from'])
        if data.get('deadline_to'):
            queryset = queryset.filter(deadline__date__lte=data['deadline_to'])
        
        # 🔹 Фильтр по клиенту
        if data.get('client'):
            queryset = queryset.filter(client=data['client'])
        
        # 🔹 Фильтр по исполнителю
        if data.get('executor'):
            queryset = queryset.filter(executor_name__icontains=data['executor'])
        
        # 🔹 Быстрые фильтры
        now = timezone.now()
        
        if data.get('urgent_only'):
            queryset = queryset.filter(
                deadline__gte=now,
                deadline__lte=now + timedelta(hours=24),
                status__in=['submitted', 'published', 'draft']
            )
        
        if data.get('overdue_only'):
            queryset = queryset.filter(
                deadline__lt=now,
                status__in=['submitted', 'published', 'draft']
            )
        
        if data.get('has_documents'):
            queryset = queryset.filter(documents__isnull=False).distinct()
        
        # 🔹 Сортировка
        sort = data.get('sort') or '-deadline'
        queryset = queryset.order_by(sort)
        
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['filter_form'] = self.get_form()
        context['total_count'] = self.get_queryset().count()
        
        now = timezone.now()
        context['now'] = now
        
        # Индикаторы срочности
        context['overdue_count'] = Tender.objects.filter(
            deadline__lt=now,
            status__in=['submitted', 'published', 'draft']
        ).count()
        
        context['urgent_count'] = Tender.objects.filter(
            deadline__gte=now,
            deadline__lte=now + timedelta(hours=24),
            status__in=['submitted', 'published', 'draft']
        ).count()
        
        # Активные фильтры для отображения бейджей
        context['active_filters'] = self._get_active_filters()
        
        return context
    
    def _get_active_filters(self):
        """Возвращает список активных фильтров для отображения"""
        form = self.get_form()
        if not form.is_valid():
            return []
        
        data = form.cleaned_data
        filters = []
        
        if data.get('search'):
            filters.append({'name': 'Поиск', 'value': data['search'], 'param': 'search'})
        
        if data.get('status'):
            status_labels = dict(TenderFilterForm.base_fields['status'].choices)
            for s in data['status']:
                filters.append({'name': 'Статус', 'value': status_labels.get(s, s), 'param': 'status', 'value_param': s})
        
        if data.get('amount_from'):
            filters.append({'name': 'Сумма от', 'value': f'{data["amount_from"]:,.0f} ₽', 'param': 'amount_from'})
        if data.get('amount_to'):
            filters.append({'name': 'Сумма до', 'value': f'{data["amount_to"]:,.0f} ₽', 'param': 'amount_to'})
        
        if data.get('deadline_from'):
            filters.append({'name': 'Дедлайн от', 'value': data['deadline_from'].strftime('%d.%m.%Y'), 'param': 'deadline_from'})
        if data.get('deadline_to'):
            filters.append({'name': 'Дедлайн до', 'value': data['deadline_to'].strftime('%d.%m.%Y'), 'param': 'deadline_to'})
        
        if data.get('client'):
            filters.append({'name': 'Клиент', 'value': str(data['client']), 'param': 'client'})
        
        if data.get('executor'):
            filters.append({'name': 'Исполнитель', 'value': data['executor'], 'param': 'executor'})
        
        if data.get('urgent_only'):
            filters.append({'name': 'Срочные', 'value': '24 часа', 'param': 'urgent_only'})
        
        if data.get('overdue_only'):
            filters.append({'name': 'Просроченные', 'value': '', 'param': 'overdue_only'})
        
        if data.get('has_documents'):
            filters.append({'name': 'С документами', 'value': '', 'param': 'has_documents'})
        
        return filters

class TenderCreateView(LoginRequiredMixin, CreateView):
    model = Tender
    form_class = TenderForm  # ← было fields = [...]
    template_name = 'tenders/tender_form.html'
    success_url = reverse_lazy('tenders:list')

    def form_valid(self, form):
        form.instance.author = self.request.user
        response = super().form_valid(form)
        
        # Логируем создание
        self.object._audit_user = self.request.user
        self.object._audit_request = self.request
        from .audit import log_action
        log_action(
            user=self.request.user,
            action='create',
            model_name='Tender',
            object_id=self.object.pk,
            object_repr=str(self.object),
            request=self.request
        )
        
        messages.success(self.request, f'✅ Тендер "{self.object.customer_name}" создан!')
        return response


class TenderUpdateView(LoginRequiredMixin, UpdateView):
    model = Tender
    form_class = TenderForm
    template_name = 'tenders/tender_form.html'
    success_url = reverse_lazy('tenders:list')
    
    def delete(self, request, *args, **kwargs):
        tender = self.get_object()
        messages.success(request, f'🗑️ Тендер "{tender.customer_name}" удалён')
        return super().delete(request, *args, **kwargs)

    def form_valid(self, form):
        if form.instance.client:
            form.instance.customer_name = form.instance.client.name
        
        response = super().form_valid(form)
        
        # Логируем обновление
        from .audit import log_action
        log_action(
            user=self.request.user,
            action='update',
            model_name='Tender',
            object_id=self.object.pk,
            object_repr=str(self.object),
            request=self.request
        )
        
        messages.success(self.request, f'✅ Тендер "{self.object.customer_name}" обновлён!')
        return response
    
class TenderDeleteView(LoginRequiredMixin, DeleteView):
    model = Tender
    template_name = 'tenders/tender_confirm_delete.html'
    success_url = reverse_lazy('tenders:list')
    
    def delete(self, request, *args, **kwargs):
        tender = self.get_object()
        tender._audit_user = request.user
        tender._audit_request = request
        response = super().delete(request, *args, **kwargs)
        messages.success(request, f'🗑️ Тендер "{tender.customer_name}" удалён')
        return response

# ============================================
# 🔹 ДАШБОРД
# ============================================
import json
from collections import defaultdict

def dashboard(request):
    """Дашборд со статистикой и графиками"""
    now = timezone.now()
    six_months_ago = now - timedelta(days=180)
    
    # 🔹 БАЗОВАЯ СТАТИСТИКА
    total_tenders = Tender.objects.count()
    won_tenders = Tender.objects.filter(status='won').count()
    active_tenders = Tender.objects.filter(status__in=['submitted', 'published']).count()
    lost_tenders = Tender.objects.filter(status='lost').count()
    draft_tenders = Tender.objects.filter(status='draft').count()
    
    # Конверсия
    conversion = round((won_tenders / total_tenders * 100), 1) if total_tenders > 0 else 0
    
    # 🔹 ФИНАНСЫ
    won_qs = Tender.objects.filter(status='won')
    totals = won_qs.aggregate(
        total_final=Sum('final_amount'),
        total_cost=Sum('cost')
    )
    total_final = totals['total_final'] or 0
    total_cost = totals['total_cost'] or 0
    profit = total_final - total_cost
    
    if total_cost > 0:
        markup_percent = round((profit / total_cost) * 100, 1)
    else:
        markup_percent = 0

    # 🔹 ГРАФИК 1: Статусы тендеров (круговая)
    status_stats = Tender.objects.values('status').annotate(count=Count('id')).order_by('status')
    status_labels = []
    status_data = []
    status_colors = {
        'draft': '#6c757d',
        'submitted': '#0dcaf0',
        'published': '#0d6efd',
        'won': '#198754',
        'lost': '#dc3545',
        'cancelled': '#adb5bd',
    }
    status_display = {
        'draft': 'Черновики',
        'submitted': 'Поданные',
        'published': 'Опубликованные',
        'won': 'Выигранные',
        'lost': 'Проигранные',
        'cancelled': 'Отменённые',
    }
    
    for item in status_stats:
        status = item['status']
        status_labels.append(status_display.get(status, status))
        status_data.append(item['count'])
    
    # 🔹 ГРАФИК 2: Тендеры по месяцам (линейный)
    tenders_qs = Tender.objects.filter(
        deadline__isnull=False,
        deadline__gte=six_months_ago
    ).values('deadline', 'status')
    
    monthly_data = defaultdict(lambda: {'total': 0, 'won': 0})
    for t in tenders_qs:
        month = t['deadline'].strftime('%Y-%m')
        monthly_data[month]['total'] += 1
        if t['status'] == 'won':
            monthly_data[month]['won'] += 1
    
    monthly_stats = sorted(monthly_data.items())
    month_labels = [m[0] for m in monthly_stats]
    month_totals = [m[1]['total'] for m in monthly_stats]
    month_won = [m[1]['won'] for m in monthly_stats]
    
    # 🔹 ГРАФИК 3: Прибыль по клиентам (столбчатый)
    client_profit = Tender.objects.filter(
        status='won',
        client__isnull=False
    ).values('client__name').annotate(
        profit_sum=Sum(F('final_amount') - F('cost'))
    ).order_by('-profit_sum')[:10]
    
    client_labels = [c['client__name'][:20] for c in client_profit]
    client_data = [float(c['profit_sum'] or 0) for c in client_profit]
    
    # 🔹 ВОРОНКА ПРОДАЖ
    funnel_data = {
        'draft': Tender.objects.filter(status='draft').count(),
        'submitted': Tender.objects.filter(status__in=['submitted', 'published']).count(),
        'won': Tender.objects.filter(status='won').count(),
    }
    
    funnel_conversion = {}
    if funnel_data['draft'] > 0:
        funnel_conversion['draft_to_submitted'] = round(
            (funnel_data['submitted'] / funnel_data['draft']) * 100, 1
        )
    else:
        funnel_conversion['draft_to_submitted'] = 0
    
    if funnel_data['submitted'] > 0:
        funnel_conversion['submitted_to_won'] = round(
            (funnel_data['won'] / funnel_data['submitted']) * 100, 1
        )
    else:
        funnel_conversion['submitted_to_won'] = 0
    
    # 🔹 Ближайшие дедлайны
    upcoming_deadlines = Tender.objects.filter(
        deadline__gte=now,
    ).order_by('deadline')[:5]
    
    # 🔹 ОТЛАДКА
    print("📊 DEBUG status_labels:", status_labels)
    print("📊 DEBUG status_data:", status_data)
    print("📊 DEBUG month_labels:", month_labels)
    print("📊 DEBUG client_labels:", client_labels)
    print("📊 DEBUG funnel_data:", funnel_data)
    
    # 🔹 КОНТЕКСТ
    context = {
        'total_tenders': total_tenders,
        'active_tenders': active_tenders,
        'won_tenders': won_tenders,
        'lost_tenders': lost_tenders,
        'draft_tenders': draft_tenders,
        'win_rate': conversion,
        'total_profit': profit,
        'avg_markup': markup_percent,
        'upcoming_deadlines': upcoming_deadlines,
        
        # Данные для графиков - как JSON строки
        'status_labels_json': json.dumps(status_labels, ensure_ascii=False),
        'status_data_json': json.dumps(status_data),
        'status_colors_json': json.dumps([status_colors.get(s['status'], '#6c757d') for s in status_stats]),
        'month_labels_json': json.dumps(month_labels),
        'month_totals_json': json.dumps(month_totals),
        'month_won_json': json.dumps(month_won),
        'client_labels_json': json.dumps(client_labels, ensure_ascii=False),
        'client_data_json': json.dumps(client_data),
        
        # Воронка продаж
        'funnel_data': funnel_data,
        'funnel_conversion': funnel_conversion,
        'funnel_data_json': json.dumps(funnel_data),
    }
    
    return render(request, 'tenders/dashboard.html', context)


    

# ============================================
# 🔹 ЭКСПОРТ В EXCEL
# ============================================
def export_tenders_excel(request):
    """Экспорт всех тендеров в Excel"""
    tenders_qs = Tender.objects.all().order_by('-deadline')
    excel_file = export_tenders_to_excel(tenders_qs)
    
    response = HttpResponse(
        excel_file,
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename=tenders_{timezone.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
    return response


def export_dashboard_excel(request):
    """Экспорт статистики дашборда в Excel"""
    excel_file = export_dashboard_stats_to_excel()
    
    response = HttpResponse(
        excel_file,
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = f'attachment; filename=dashboard_stats_{timezone.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
    return response


# ============================================
# 🔹 ПОИСК ТЕНДЕРОВ ОНЛАЙН
# ============================================
class ExternalSearchView(LoginRequiredMixin, View):
    """Поиск тендеров на zakupki.gov.ru"""
    
    def get(self, request):
        form = SearchTenderForm()
        return render(request, 'tenders/external_search.html', {'form': form})
    
    def post(self, request):
        form = SearchTenderForm(request.POST)
        
        if form.is_valid():
            keyword = form.cleaned_data['keyword']
            region = form.cleaned_data['region'] or None
            max_results = form.cleaned_data['max_results']
            
            tenders = search_tenders_on_zakupki(keyword, region, max_results)
        
            return render(request, 'tenders/external_search.html', {
                'form': form,
                'tenders': tenders,
                'search_performed': True,
                'keyword': keyword,
            })
        return render(request, 'tenders/external_search.html', {'form': form})





class ImportTenderView(LoginRequiredMixin, View):
    """Импорт тендера с zakupki.gov.ru в нашу систему"""
    
    def get(self, request):
        """Простая страница-заглушка"""
        messages.info(request, 'ℹ️ Используйте страницу поиска для импорта тендеров')
        return redirect('tenders:external_search')
    
    def post(self, request):
        """Обработка импорта тендера (AJAX)"""
        try:
            customer_name = request.POST.get('customer_name', 'Не указан')[:200]
            
            # 🔹 АВТОПОИСК клиента по названию заказчика
            client = None
            if customer_name and customer_name != 'Не указан':
                client = Client.objects.filter(
                    Q(name__iexact=customer_name) | 
                    Q(name__icontains=customer_name)
                ).first()
            
            tender = Tender(
                customer_name=customer_name,
                client=client,
                initial_amount=request.POST.get('initial_amount') or 0,
                deadline=request.POST.get('deadline') or None,
                status='draft',
                executor_name='',
                procedure_url=request.POST.get('source_url', ''),
                author=request.user,
                source_url=request.POST.get('source_url', ''),
                external_id=request.POST.get('external_id', ''),
                comment=f"Импортирован с zakupki.gov.ru: {request.POST.get('title', '')}"[:1000],
            )
            tender.save()
            
            # 🔹 Формируем ответ для AJAX
            if client:
                message = f'✅ Тендер "{customer_name}" импортирован и привязан к клиенту "{client.name}"!'
                status_type = 'success'
            else:
                message = f'⚠️ Тендер "{customer_name}" импортирован. Клиент не найден.'
                status_type = 'warning'
            
            return JsonResponse({
                'success': True,
                'message': message,
                'status': status_type,
                'tender_id': tender.pk,
                'tender_name': customer_name,
            })
            
        except Exception as e:
            return JsonResponse({
                'success': False,
                'message': f'❌ Ошибка импорта: {e}',
                'status': 'error',
            }, status=500)


# ============================================
# 🔹 CRM: КЛИЕНТЫ
# ============================================
class ClientListView(LoginRequiredMixin, ListView):
    #Список клиентов"""
    model = Client
    template_name = 'tenders/client_list.html'
    context_object_name = 'clients'
    paginate_by = 20
    
    def get_queryset(self):
        queryset = Client.objects.select_related('manager', 'created_by')
        
        # Фильтрация по поиску
        search = self.request.GET.get('search')
        if search:
            queryset = queryset.filter(
                db_models.Q(name__icontains=search) |
                db_models.Q(inn__icontains=search) |
                db_models.Q(email__icontains=search) |
                db_models.Q(contact_person__icontains=search)
            )
        
        # Фильтрация по статусу
        status = self.request.GET.get('status')
        if status:
            queryset = queryset.filter(status=status)
        
        # Фильтрация по менеджеру
        manager = self.request.GET.get('manager')
        if manager:
            queryset = queryset.filter(manager_id=manager)
        
        return queryset
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['now'] = timezone.now() 
        
        # 🔹 Индикаторы срочности
        now = timezone.now()
        
        # Количество просроченных тендеров
        context['overdue_count'] = Tender.objects.filter(
            deadline__lt=now,
            status__in=['submitted', 'published', 'draft']
        ).count()
        
        # Количество срочных (в ближайшие 24 часа)
        context['urgent_count'] = Tender.objects.filter(
            deadline__gte=now,
            deadline__lte=now + timedelta(hours=24),
            status__in=['submitted', 'published', 'draft']
        ).count()
        
        # Количество с дедлайном в ближайшие 3 дня
        context['upcoming_count'] = Tender.objects.filter(
            deadline__gte=now,
            deadline__lte=now + timedelta(days=3),
            status__in=['submitted', 'published', 'draft']
        ).count()
        
        return context

class ClientCreateView(LoginRequiredMixin, CreateView):
    #Создание клиента"""
    model = Client
    form_class = ClientForm
    template_name = 'tenders/client_form.html'
    success_url = reverse_lazy('tenders:client_list')
    
    def form_valid(self, form):
        form.instance.created_by = self.request.user
        response = super().form_valid(form)
        messages.success(self.request, f'✅ Клиент "{self.object.name}" успешно добавлен!')
        return response


class ClientUpdateView(LoginRequiredMixin, UpdateView):
    #Редактирование клиента"""
    model = Client
    form_class = ClientForm
    template_name = 'tenders/client_form.html'
    success_url = reverse_lazy('tenders:client_list')
    
    def form_valid(self, form):
        response = super().form_valid(form)
        messages.success(self.request, f'✅ Данные клиента "{self.object.name}" обновлены!')
        return response


class ClientDetailView(LoginRequiredMixin, DetailView):
    #Детальная информация о клиенте"""
    model = Client
    template_name = 'tenders/client_detail.html'
    context_object_name = 'client'
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['now'] = timezone.now()
        # 🔹 Здесь self.object — это текущий клиент (правильно!)
        related_tenders = Tender.objects.filter(client=self.object).order_by('-deadline')
        
        # Статистика по клиенту
        total_tenders = related_tenders.count()
        won_tenders = related_tenders.filter(status='won').count()
        total_profit = related_tenders.filter(status='won').aggregate(
            total=Sum(F('final_amount') - F('cost'))
        )['total'] or 0
        
        context['related_tenders'] = related_tenders[:10]
        context['client_stats'] = {
            'total_tenders': total_tenders,
            'won_tenders': won_tenders,
            'win_rate': round(won_tenders / total_tenders * 100, 1) if total_tenders > 0 else 0,
            'total_profit': total_profit,
        }
        
        return context

class ClientDeleteView(LoginRequiredMixin, DeleteView):
    """Удаление клиента"""
    model = Client
    template_name = 'tenders/client_confirm_delete.html'
    success_url = reverse_lazy('tenders:client_list')
    
    def delete(self, request, *args, **kwargs):
        client = self.get_object()
        messages.success(request, f'🗑️ Клиент "{client.name}" удалён')
        return super().delete(request, *args, **kwargs)
    



class TenderDocumentUploadView(LoginRequiredMixin, View):
    """Загрузка документа к тендеру"""
    
    def post(self, request, pk):
        tender = Tender.objects.get(pk=pk)
        
        # Проверка прав доступа
        if tender.author != request.user:
            messages.error(request, '❌ У вас нет прав для добавления документов')
            return redirect('tenders:list')
        
        form = TenderDocumentForm(request.POST, request.FILES)
        if form.is_valid():
            document = form.save(commit=False)
            document.tender = tender
            document.uploaded_by = request.user
            document.save()
            messages.success(request, f'📎 Документ "{document.name}" загружен!')
        else:
            for error in form.errors.values():
                messages.error(request, error)
        
        return redirect('tenders:update', pk=pk)


class TenderDocumentDeleteView(LoginRequiredMixin, View):
    """Удаление документа"""
    
    def post(self, request, pk):
        document = TenderDocument.objects.get(pk=pk)
        
        # Проверка прав доступа
        if document.tender.author != request.user:
            messages.error(request, '❌ У вас нет прав для удаления')
            return redirect('tenders:list')
        
        document_name = document.name
        document.file.delete()  # Удаляем файл с диска
        document.delete()
        messages.success(request, f'🗑️ Документ "{document_name}" удалён')
        
        return redirect('tenders:update', pk=document.tender.pk)
    
   


class AuditLogView(LoginRequiredMixin, ListView):
    """Просмотр лога аудита"""
    model = AuditLog
    template_name = 'tenders/audit_log.html'
    context_object_name = 'logs'
    paginate_by = 50
    
    def get_queryset(self):
        queryset = AuditLog.objects.select_related('user')
        
        # Фильтрация по пользователю
        user = self.request.GET.get('user')
        if user:
            queryset = queryset.filter(user_id=user)
        
        # Фильтрация по действию
        action = self.request.GET.get('action')
        if action:
            queryset = queryset.filter(action=action)
        
        # Фильтрация по модели
        model_name = self.request.GET.get('model')
        if model_name:
            queryset = queryset.filter(model_name=model_name)
        
        # Фильтрация по дате
        date_from = self.request.GET.get('date_from')
        if date_from:
            from datetime import datetime
            try:
                date_from = datetime.strptime(date_from, '%Y-%m-%d')
                queryset = queryset.filter(timestamp__gte=date_from)
            except:
                pass
        
        return queryset
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['users'] = User.objects.filter(is_active=True)
        context['actions'] = AuditLog.ACTION_CHOICES
        context['models'] = ['Tender', 'Client', 'AuditLog']
        return context
    
class TenderKanbanView(LoginRequiredMixin, View):
    """Канбан-доска для визуального управления тендерами"""
    template_name = 'tenders/kanban.html'
    
    def get(self, request):
        tenders = Tender.objects.all().order_by('-deadline')
        
        # Группируем тендеры по статусам
        context = {
            'draft': tenders.filter(status='draft'),
            'process': tenders.filter(status__in=['submitted', 'published']),
            'won': tenders.filter(status='won'),
            'lost': tenders.filter(status__in=['lost', 'cancelled']),
        }
        
        return render(request, self.template_name, context)

class TenderStatusUpdateView(LoginRequiredMixin, View):
    def post(self, request):
        try:
            import json
            data = json.loads(request.body)
            
            tender_id = data.get('tender_id')
            new_status = data.get('status')
            
            if not tender_id or not new_status:
                return JsonResponse({
                    'status': 'error',
                    'message': f'Missing data: tender_id={tender_id}, status={new_status}'
                }, status=400)
            
            tender = Tender.objects.get(pk=tender_id)
            old_status = tender.status
            
            # Маппинг статусов из колонки в значение модели
            status_map = {
                'draft': 'draft',
                'process': 'submitted',  # process -> submitted
                'won': 'won',
                'lost': 'lost'
            }
            
            actual_status = status_map.get(new_status, new_status)
            tender.status = actual_status
            tender.save()
            
            return JsonResponse({'status': 'success'})
            
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


class TenderCalendarView(LoginRequiredMixin, View):
    """Интерактивный календарь дедлайнов"""
    template_name = 'tenders/calendar.html'
    
    def get(self, request):
        return render(request, self.template_name)


class TenderCalendarDataView(LoginRequiredMixin, View):
    """API endpoint для получения данных календаря"""
    
    def get(self, request):
        # Получаем параметры фильтрации
        start = request.GET.get('start')
        end = request.GET.get('end')
        
        tenders = Tender.objects.filter(deadline__isnull=False)
        
        # Фильтрация по исполнителю
        executor = request.GET.get('executor')
        if executor:
            tenders = tenders.filter(executor_name__icontains=executor)
        
        # Фильтрация по статусу
        status = request.GET.get('status')
        if status:
            tenders = tenders.filter(status=status)
        
        # Формируем данные для календаря
        events = []
        for tender in tenders:
            # Цвет в зависимости от статуса
            color_map = {
                'draft': '#6c757d',      # серый
                'submitted': '#0dcaf0',  # голубой
                'published': '#0d6efd',  # синий
                'won': '#198754',        # зелёный
                'lost': '#dc3545',       # красный
                'cancelled': '#adb5bd',  # светло-серый
            }
            
            color = color_map.get(tender.status, '#6c757d')
            
            events.append({
                'id': tender.pk,
                'title': f"{tender.customer_name[:30]} ({tender.initial_amount:,.0f}₽)",
                'start': tender.deadline.isoformat(),
                'url': f"/tenders/{tender.pk}/edit/",
                'color': color,
                'extendedProps': {
                    'status': tender.get_status_display(),
                    'amount': tender.initial_amount,
                    'executor': tender.executor_name or 'Не указан'
                }
            })
        
        return JsonResponse(events, safe=False)


@login_required
def add_comment(request, pk):
    """Добавление комментария к тендеру (AJAX)"""
    tender = get_object_or_404(Tender, pk=pk)
    
    if request.method == 'POST':
        form = CommentForm(request.POST)
        if form.is_valid():
            comment = form.save(commit=False)
            comment.tender = tender
            comment.author = request.user
            comment.save()
            
            # Создаём запись в audit log
            AuditLog.objects.create(
                user=request.user,
                action='create',
                model_name='Comment',
                object_id=comment.id,
                object_repr=f'Комментарий к {tender.customer_name}',
                ip_address=get_client_ip(request),
                user_agent=request.META.get('HTTP_USER_AGENT', '')
            )
            
            # Возвращаем HTML для нового комментария
            return JsonResponse({
                'success': True,
                'html': render_to_string('tenders/comment_item.html', {
                    'comment': comment,
                    'user': request.user
                }, request=request)
            })
        else:
            return JsonResponse({'success': False, 'errors': form.errors})
    
    return JsonResponse({'success': False, 'error': 'Метод не разрешён'})


@login_required
def edit_comment(request, pk):
    """Редактирование комментария"""
    comment = get_object_or_404(Comment, pk=pk)
    
    # Проверяем, что пользователь является автором
    if comment.author != request.user:
        return JsonResponse({'success': False, 'error': 'Нет прав'})
    
    if request.method == 'POST':
        form = CommentForm(request.POST, instance=comment)
        if form.is_valid():
            form.save()
            return JsonResponse({
                'success': True,
                'html': render_to_string('tenders/comment_item.html', {
                    'comment': comment,
                    'user': request.user
                }, request=request)
            })
        else:
            return JsonResponse({'success': False, 'errors': form.errors})
    
    return JsonResponse({'success': False, 'error': 'Метод не разрешён'})


@login_required
def delete_comment(request, pk):
    """Удаление комментария"""
    comment = get_object_or_404(Comment, pk=pk)
    
    # Проверяем, что пользователь является автором
    if comment.author != request.user:
        return JsonResponse({'success': False, 'error': 'Нет прав'})
    
    if request.method == 'POST':
        tender = comment.tender
        comment.delete()
        
        AuditLog.objects.create(
            user=request.user,
            action='delete',
            model_name='Comment',
            object_repr=f'Удалён комментарий к {tender.customer_name}',
            ip_address=get_client_ip(request),
            user_agent=request.META.get('HTTP_USER_AGENT', '')
        )
        
        return JsonResponse({'success': True})
    
    return JsonResponse({'success': False, 'error': 'Метод не разрешён'})

@login_required
def export_tenders_pdf(request):
    """Экспорт списка тендеров в PDF"""
    from .pdf_utils import render_tenders_list_pdf
    from django.utils import timezone
    from django.db.models import Sum
    
    tenders = Tender.objects.filter(author=request.user).select_related('author', 'client')
    
    status = request.GET.get('status')
    if status:
        tenders = tenders.filter(status=status)
    
    total_count = tenders.count()
    total_initial = tenders.aggregate(total=Sum('initial_amount'))['total'] or 0
    won_count = tenders.filter(status='won').count()
    
    context = {
        'current_date': timezone.now().strftime('%d.%m.%Y %H:%M'),
        'total_count': total_count,
        'total_initial': total_initial,
        'won_count': won_count,
    }
    
    return render_tenders_list_pdf(tenders, context)


@login_required
def export_dashboard_pdf(request):
    """Экспорт дашборда в PDF"""
    from .pdf_utils import render_dashboard_pdf
    from django.utils import timezone
    from django.db.models import Sum, F, Count, Q
    
    user = request.user
    tenders = Tender.objects.filter(author=user)
    
    # Общая статистика
    total_tenders = tenders.count()
    active_tenders = tenders.filter(status__in=['draft', 'submitted']).count()
    won_tenders = tenders.filter(status='won').count()
    lost_tenders = tenders.filter(status='lost').count()
    win_rate = round(won_tenders / total_tenders * 100, 1) if total_tenders > 0 else 0
    
    # Финансовые показатели
    totals = tenders.aggregate(
        total_initial=Sum('initial_amount'),
        total_won=Sum('final_amount', filter=Q(status='won')),
        total_cost=Sum('cost', filter=Q(status='won'))
    )
    
    total_initial_amount = totals['total_initial'] or 0
    total_won_amount = totals['total_won'] or 0
    total_cost = totals['total_cost'] or 0
    total_profit = total_won_amount - total_cost
    average_markup = (total_profit / total_cost * 100) if total_cost else 0
    
    # Статистика по статусам
    status_stats = []
    for code, name in Tender._meta.get_field('status').choices:
        count = tenders.filter(status=code).count()
        percent = round(count / total_tenders * 100, 1) if total_tenders > 0 else 0
        status_stats.append({
            'code': code,
            'name': name,
            'count': count,
            'percent': percent
        })
    
    context = {
        'current_date': timezone.now().strftime('%d.%m.%Y %H:%M'),
        'total_tenders': total_tenders,
        'active_tenders': active_tenders,
        'won_tenders': won_tenders,
        'lost_tenders': lost_tenders,
        'win_rate': win_rate,
        'total_initial_amount': total_initial_amount,
        'total_won_amount': total_won_amount,
        'total_cost': total_cost,
        'total_profit': total_profit,
        'average_markup': average_markup,
        'status_stats': status_stats,
    }
    
    return render_dashboard_pdf(context)

@login_required
def export_tender_detail_pdf(request, pk):
    """Экспорт карточки тендера в PDF"""
    from .pdf_utils import render_tender_detail_pdf
    from django.utils import timezone
    
    tender = get_object_or_404(Tender, pk=pk, author=request.user)
    
    context = {
        'current_date': timezone.now().strftime('%d.%m.%Y %H:%M'),
    }
    
    return render_tender_detail_pdf(tender, context)

class TenderplanSearchView(LoginRequiredMixin, View):
    """Поиск тендеров на tenderplan.ru"""
    template_name = 'tenders/tenderplan_search.html'
    
    def get(self, request):
        form = TenderplanSearchForm()
        return render(request, self.template_name, {
            'form': form,
            'search_performed': False,
        })
    
    def post(self, request):
        form = TenderplanSearchForm(request.POST)
        tenders = []
        error = None
        
        if form.is_valid():
            keyword = form.cleaned_data['keyword']
            region = form.cleaned_data.get('region')
            max_results = form.cleaned_data['max_results']
            
            try:
                tenders = api_client.search_tenders(keyword, region, max_results)
            except Exception as e:
                error = str(e)
        
        return render(request, self.template_name, {
            'form': form,
            'tenders': tenders,
            'search_performed': True,
            'error': error,
            'keyword': form.cleaned_data.get('keyword', '') if form.is_valid() else '',
        })


class ImportFromTenderplanView(LoginRequiredMixin, View):
    """Импорт тендера с tenderplan.ru"""
    
    def post(self, request):
        try:
            external_id = request.POST.get('external_id', '')
            customer_name = request.POST.get('customer_name', 'Не указан')[:200]

            # Конвертируем сумму: заменяем запятую на точку
            amount_str = request.POST.get('initial_amount', '0') or '0'
            try:
                # Заменяем запятую на точку и удаляем пробелы
                amount_str = amount_str.replace(',', '.').replace(' ', '').strip()
                initial_amount = float(amount_str)
            except (ValueError, TypeError):
                initial_amount = 0.0
        
            deadline = request.POST.get('deadline') or None
            procedure_url = request.POST.get('procedure_url', '')
            title = request.POST.get('title', '')[:200]
            
            # Автопоиск клиента
            client = None
            if customer_name and customer_name != 'Не указан':
                client = Client.objects.filter(
                    Q(name__iexact=customer_name) |
                    Q(name__icontains=customer_name)
                ).first()
            
            # Парсим дату
            deadline_dt = None
            if deadline:
                from datetime import datetime
                for fmt in ['%Y-%m-%d %H:%M:%S', '%Y-%m-%d', '%d.%m.%Y']:
                    try:
                        deadline_dt = datetime.strptime(deadline, fmt)
                        break
                    except:
                        continue
            
            tender = Tender(
                customer_name=customer_name,
                client=client,
                initial_amount=initial_amount,
                deadline=deadline_dt,
                status='draft',
                executor_name='',
                procedure_url=procedure_url,
                author=request.user,
                source_url=procedure_url,
                external_id=external_id,
                comment=f"Импортирован с tenderplan.ru: {title}"[:1000],
            )
            tender.save()
            
            if client:
                message = f'✅ Тендер "{customer_name}" импортирован и привязан к клиенту "{client.name}"!'
                status_type = 'success'
            else:
                message = f'⚠️ Тендер "{customer_name}" импортирован. Клиент не найден.'
                status_type = 'warning'
            
            return JsonResponse({
                'success': True,
                'message': message,
                'status': status_type,
                'tender_id': tender.pk,
            })
        
        except Exception as e:
            return JsonResponse({
                'success': False,
                'message': f'❌ Ошибка импорта: {e}',
                'status': 'error',
            }, status=500)

def post(self, request):
    form = TenderplanSearchForm(request.POST)
    tenders = []
    error = None
    
    if form.is_valid():
        keyword = form.cleaned_data['keyword']
        region = form.cleaned_data.get('region')
        max_results = form.cleaned_data['max_results']
        
        print(f"🔍 Поиск: keyword={keyword}, region={region}, max={max_results}")
        
        try:
            tenders = api_client.search_tenders(keyword, region, max_results)
            print(f"✅ Найдено: {len(tenders)}")
        except Exception as e:
            error = str(e)
            print(f"❌ Ошибка: {error}")
    
    return render(request, self.template_name, {
        'form': form,
        'tenders': tenders,
        'search_performed': True,
        'error': error,
        'keyword': form.cleaned_data.get('keyword', '') if form.is_valid() else '',
    })





def export_kanban_to_word(request):
    """Экспорт канбан-доски в Word"""
    # Создаём документ
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('📋 Канбан-доска SRM Тендеры', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Дата формирования
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f'Дата формирования: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    run.italic = True
    
    # Получаем все тендеры по статусам
    draft = Tender.objects.filter(status='draft')
    process = Tender.objects.filter(status__in=['submitted', 'published'])
    won = Tender.objects.filter(status='won')
    lost = Tender.objects.filter(status__in=['lost', 'cancelled'])
    
    # Функция добавления колонки
    def add_column(title, tenders, color):
        doc.add_heading(title, level=1)
        
        if not tenders:
            p = doc.add_paragraph('Нет тендеров')
            p.runs[0].italic = True
            return
        
        for tender in tenders:
            # Создаём таблицу для каждой карточки
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid Accent 1'
            
            # Заголовок карточки
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = f'#{tender.id} - {tender.customer_name[:50]}'
            hdr_cells[0].paragraphs[0].runs[0].bold = True
            
            # Объединяем ячейки для заголовка
            if len(table.columns) > 1:
                table.rows[0].cells[1].merge(table.rows[0].cells[1])
            
            # Данные
            data = [
                ('Заказчик', tender.customer_name),
                ('Сумма', f'{tender.initial_amount:,.0f} ₽'),
                ('Дедлайн', tender.deadline.strftime('%d.%m.%Y %H:%M') if tender.deadline else 'Не указан'),
                ('Исполнитель', tender.executor_name or 'Не назначен'),
                ('Статус', tender.get_status_display()),
            ]
            
            if tender.client:
                data.append(('Клиент', tender.client.name))
            
            for label, value in data:
                row_cells = table.add_row().cells
                row_cells[0].text = label
                row_cells[1].text = value
                row_cells[0].paragraphs[0].runs[0].bold = True
                row_cells[0].width = Inches(1.5)
            
            # Пустая строка между карточками
            doc.add_paragraph()
    
    # Добавляем колонки
    add_column('📝 Черновики', draft, 'gray')
    add_column('📤 В работе', process, 'blue')
    add_column('🏆 Выигранные', won, 'green')
    add_column('❌ Закрыты', lost, 'red')
    
    # Итоговая статистика
    doc.add_page_break()
    doc.add_heading('📊 Итоговая статистика', level=1)
    
    stats_table = doc.add_table(rows=1, cols=3)
    stats_table.style = 'Medium Grid 1 Accent 1'
    hdr_cells = stats_table.rows[0].cells
    hdr_cells[0].text = 'Категория'
    hdr_cells[1].text = 'Количество'
    hdr_cells[2].text = 'Общая сумма'
    
    categories = [
        ('Черновики', draft),
        ('В работе', process),
        ('Выигранные', won),
        ('Закрыты', lost),
    ]
    
    total_sum = 0
    for name, qs in categories:
        count = qs.count()
        amount = sum(t.initial_amount or 0 for t in qs)
        total_sum += amount
        
        row_cells = stats_table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = str(count)
        row_cells[2].text = f'{amount:,.0f} ₽'
    
    # Итого
    total_row = stats_table.add_row().cells
    total_row[0].text = 'ВСЕГО'
    total_row[0].paragraphs[0].runs[0].bold = True
    total_row[1].text = str(Tender.objects.count())
    total_row[1].paragraphs[0].runs[0].bold = True
    total_row[2].text = f'{total_sum:,.0f} ₽'
    total_row[2].paragraphs[0].runs[0].bold = True
    
    # Сохраняем в response
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename=kanban_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    
    doc.save(response)
    return response




def export_dashboard_to_word(request):
    """Экспорт дашборда в Word"""
    doc = Document()
    
    # Ориентация ландшафтная для таблиц
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    
    # Заголовок
    title = doc.add_heading('📊 Дашборд SRM Тендеры', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Дата и пользователь
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f'Дата: {datetime.now().strftime("%d.%m.%Y %H:%M")}\n')
    run.italic = True
    if request.user.is_authenticated:
        run = p.add_run(f'Пользователь: {request.user.username}')
        run.italic = True
    
    # ============================================
    # 🔹 ОСНОВНАЯ СТАТИСТИКА
    # ============================================
    doc.add_heading('📈 Основная статистика', level=1)
    
    now = timezone.now()
    total_tenders = Tender.objects.count()
    won_tenders = Tender.objects.filter(status='won').count()
    active_tenders = Tender.objects.filter(status__in=['submitted', 'published']).count()
    lost_tenders = Tender.objects.filter(status='lost').count()
    draft_tenders = Tender.objects.filter(status='draft').count()
    
    # Конверсия
    conversion = round((won_tenders / total_tenders * 100), 1) if total_tenders > 0 else 0
    
    # Таблица основной статистики
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Medium Grid 1 Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Показатель'
    hdr_cells[1].text = 'Значение'
    
    stats = [
        ('Всего тендеров', total_tenders),
        ('Активных', active_tenders),
        ('Выигранных', won_tenders),
        ('Проигранных', lost_tenders),
        ('Черновики', draft_tenders),
        ('Конверсия, %', f'{conversion}%'),
    ]
    
    for name, value in stats:
        row_cells = table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = str(value)
    
    doc.add_paragraph()
    
    # ============================================
    # 🔹 ФИНАНСОВЫЕ ПОКАЗАТЕЛИ
    # ============================================
    doc.add_heading('💰 Финансовые показатели', level=1)
    
    won_qs = Tender.objects.filter(status='won')
    totals = won_qs.aggregate(
        total_final=Sum('final_amount'),
        total_cost=Sum('cost')
    )
    total_final = totals['total_final'] or 0
    total_cost = totals['total_cost'] or 0
    profit = total_final - total_cost
    
    if total_cost > 0:
        markup_percent = round((profit / total_cost) * 100, 1)
    else:
        markup_percent = 0
    
    # Таблица финансов
    finance_table = doc.add_table(rows=1, cols=2)
    finance_table.style = 'Medium Grid 1 Accent 1'
    
    hdr_cells = finance_table.rows[0].cells
    hdr_cells[0].text = 'Показатель'
    hdr_cells[1].text = 'Сумма'
    
    finances = [
        ('Общая сумма выигранных тендеров', f'{total_final:,.0f} ₽'),
        ('Общая стоимость', f'{total_cost:,.0f} ₽'),
        ('Прибыль', f'{profit:,.0f} ₽'),
        ('Наценка, %', f'{markup_percent}%'),
    ]
    
    for name, value in finances:
        row_cells = finance_table.add_row().cells
        row_cells[0].text = name
        row_cells[1].text = value
        if 'Прибыль' in name and profit > 0:
            row_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
    
    doc.add_paragraph()
    
    # ============================================
    # 🔹 СТАТИСТИКА ПО СТАТУСАМ
    # ============================================
    doc.add_heading('📊 Распределение по статусам', level=1)
    
    status_table = doc.add_table(rows=1, cols=3)
    status_table.style = 'Light Grid Accent 1'
    
    hdr_cells = status_table.rows[0].cells
    hdr_cells[0].text = 'Статус'
    hdr_cells[1].text = 'Количество'
    hdr_cells[2].text = 'Процент'
    
    status_stats = Tender.objects.values('status').annotate(
        count=Count('id')
    ).order_by('status')
    
    status_display = {
        'draft': 'Черновик',
        'submitted': 'Подан',
        'published': 'Опубликован',
        'won': 'Выигран',
        'lost': 'Проигран',
        'cancelled': 'Отменён',
    }
    
    for item in status_stats:
        status = item['status']
        count = item['count']
        percent = round((count / total_tenders * 100), 1) if total_tenders > 0 else 0
        
        row_cells = status_table.add_row().cells
        row_cells[0].text = status_display.get(status, status)
        row_cells[1].text = str(count)
        row_cells[2].text = f'{percent}%'
    
    doc.add_paragraph()
    
    # ============================================
    # 🔹 ТОП КЛИЕНТОВ ПО ПРИБЫЛИ
    # ============================================
    doc.add_heading('🏆 Топ клиентов по прибыли', level=1)
    
    client_profit = Tender.objects.filter(
        status='won',
        client__isnull=False
    ).values('client__name').annotate(
        profit_sum=Sum(F('final_amount') - F('cost')),
        tender_count=Count('id')
    ).order_by('-profit_sum')[:10]
    
    if client_profit:
        client_table = doc.add_table(rows=1, cols=3)
        client_table.style = 'Light Grid Accent 1'
        
        hdr_cells = client_table.rows[0].cells
        hdr_cells[0].text = 'Клиент'
        hdr_cells[1].text = 'Тендеров'
        hdr_cells[2].text = 'Прибыль'
        
        for client_data in client_profit:
            row_cells = client_table.add_row().cells
            row_cells[0].text = client_data['client__name'] or 'Не указан'
            row_cells[1].text = str(client_data['tender_count'])
            row_cells[2].text = f"{client_data['profit_sum']:,.0f} ₽" if client_data['profit_sum'] else "0 ₽"
    else:
        p = doc.add_paragraph('Нет данных по клиентам')
        p.runs[0].italic = True
    
    doc.add_paragraph()
    
    # ============================================
    # 🔹 БЛИЖАЙШИЕ ДЕДЛАЙНЫ
    # ============================================
    doc.add_heading('⏰ Ближайшие дедлайны (7 дней)', level=1)
    
    seven_days_later = now + timedelta(days=7)
    upcoming = Tender.objects.filter(
        deadline__gte=now,
        deadline__lte=seven_days_later,
        status__in=['draft', 'submitted', 'published']
    ).order_by('deadline')[:15]
    
    if upcoming:
        deadline_table = doc.add_table(rows=1, cols=4)
        deadline_table.style = 'Light Grid Accent 1'
        
        hdr_cells = deadline_table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Заказчик'
        hdr_cells[2].text = 'Дедлайн'
        hdr_cells[3].text = 'Сумма'
        
        for tender in upcoming:
            row_cells = deadline_table.add_row().cells
            row_cells[0].text = f'#{tender.id}'
            row_cells[1].text = tender.customer_name[:40] if tender.customer_name else ''
            row_cells[2].text = tender.deadline.strftime('%d.%m.%Y %H:%M') if tender.deadline else ''
            row_cells[3].text = f'{tender.initial_amount:,.0f} ₽' if tender.initial_amount else ''
            
            # Подсветка срочных
            if tender.deadline and (tender.deadline - now).days <= 1:
                for cell in row_cells:
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                    cell.paragraphs[0].runs[0].bold = True
    else:
        p = doc.add_paragraph('Нет предстоящих дедлайнов')
        p.runs[0].italic = True
    
    doc.add_paragraph()
    # ============================================
    # 🔹 ПОСЛЕДНИЕ 10 ТЕНДЕРОВ
    # ============================================
    doc.add_heading('📋 Последние тендеры', level=1)
    
    recent_tenders = Tender.objects.all().order_by('-id')[:10]
    
    if recent_tenders:
        recent_table = doc.add_table(rows=1, cols=5)
        recent_table.style = 'Light Grid Accent 1'
        
        hdr_cells = recent_table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Заказчик'
        hdr_cells[2].text = 'Статус'
        hdr_cells[3].text = 'Сумма'
        hdr_cells[4].text = 'Дедлайн'
        
        for tender in recent_tenders:
            row_cells = recent_table.add_row().cells
            row_cells[0].text = f'#{tender.id}'
            row_cells[1].text = tender.customer_name[:30] if tender.customer_name else ''
            row_cells[2].text = tender.get_status_display()
            row_cells[3].text = f'{tender.initial_amount:,.0f} ₽' if tender.initial_amount else ''
            row_cells[4].text = tender.deadline.strftime('%d.%m.%Y') if tender.deadline else 'Не указан'
    
    # ============================================
    # 🔹 СВОДКА
    # ============================================
    doc.add_page_break()
    doc.add_heading('📌 Итоговая сводка', level=1)
    
    summary = f"""
Общее количество тендеров: {total_tenders}
Активных тендеров: {active_tenders}
Выиграно: {won_tenders} ({conversion}%)
Проиграно: {lost_tenders}

Финансовые результаты:
- Общая сумма выигранных тендеров: {total_final:,.0f} ₽
- Общая стоимость: {total_cost:,.0f} ₽
- Прибыль: {profit:,.0f} ₽
- Наценка: {markup_percent}%

Просроченных тендеров: {Tender.objects.filter(deadline__lt=now, status__in=['draft', 'submitted', 'published']).count()}
Срочных (24 часа): {Tender.objects.filter(deadline__gte=now, deadline__lte=now + timedelta(hours=24), status__in=['draft', 'submitted', 'published']).count()}
"""
    
    p = doc.add_paragraph(summary)
    p.runs[0].font.size = Pt(11)
    
    # Сохранение
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename=dashboard_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    
    doc.save(response)
    return response



def is_admin(user):
    return user.is_superuser

@user_passes_test(is_admin)
def user_management(request):
    """Управление пользователями (только для суперпользователя)"""
    users = User.objects.all().order_by('-date_joined')
    
    if request.method == 'POST':
        action = request.POST.get('action')
        user_id = request.POST.get('user_id')
        
        if user_id:
            user = User.objects.get(id=user_id)
            
            if action == 'reset_password':
                # Сброс пароля на случайный
                new_password = User.objects.make_random_password(length=12)
                user.set_password(new_password)
                user.save()
                messages.success(request, f'Пароль пользователя {user.username} сброшен. Новый пароль: {new_password}')
            
            elif action == 'activate':
                user.is_active = True
                user.save()
                messages.success(request, f'Пользователь {user.username} активирован')
            
            elif action == 'deactivate':
                user.is_active = False
                user.save()
                messages.success(request, f'Пользователь {user.username} деактивирован')
    
    return render(request, 'tenders/user_management.html', {'users': users})


    # ============================================
# 🔹 API ДЛЯ PUSH-УВЕДОМЛЕНИЙ
# ============================================
class OverdueTendersView(View):
    """API: Просроченные тендеры"""
    
    def get(self, request):
        overdue_count = Tender.objects.filter(
            deadline__lt=timezone.now(),
            status__in=['draft', 'submitted', 'published']
        ).count()
        
        return JsonResponse({
            'count': overdue_count,
            'timestamp': timezone.now().isoformat()
        })


class UpcomingDeadlinesView(View):
    """API: Приближающиеся дедлайны"""
    
    def get(self, request):
        three_days_later = timezone.now() + timedelta(days=3)
        upcoming_count = Tender.objects.filter(
            deadline__gte=timezone.now(),
            deadline__lte=three_days_later,
            status__in=['draft', 'submitted', 'published']
        ).count()
        
        return JsonResponse({
            'count': upcoming_count,
            'timestamp': timezone.now().isoformat()
        })


# ============================================
# 🔹 DRF VIEWSETS ДЛЯ SWAGGER
# ============================================
class TenderViewSet(viewsets.ModelViewSet):
    """API для управления тендерами"""
    queryset = Tender.objects.all()
    serializer_class = TenderSerializer
    
    def get_queryset(self):
        return Tender.objects.filter(author=self.request.user)


class ClientViewSet(viewsets.ModelViewSet):
    """API для управления клиентами"""
    queryset = Client.objects.all()
    serializer_class = ClientSerializer


# ============================================
# 🔹 КАСТОМНАЯ АДМИНКА (УПРАВЛЕНИЕ ПОЛЬЗОВАТЕЛЯМИ)
# ============================================
from django.contrib.auth.models import User
from django.contrib.auth.decorators import user_passes_test

def is_admin(user):
    return user.is_superuser

@user_passes_test(is_admin)
def user_management(request):
    """Управление пользователями (только для суперпользователя)"""
    users = User.objects.all().order_by('-date_joined')
    
    if request.method == 'POST':
        action = request.POST.get('action')
        user_id = request.POST.get('user_id')
        
        if user_id:
            user = User.objects.get(id=user_id)
            
            if action == 'reset_password':
                from django.utils.crypto import get_random_string
                new_password = get_random_string(length=12)
                user.set_password(new_password)
                user.save()
                messages.success(request, f'Пароль пользователя {user.username} сброшен. Новый пароль: {new_password}')
            
            elif action == 'activate':
                user.is_active = True
                user.save()
                messages.success(request, f'Пользователь {user.username} активирован')
            
            elif action == 'deactivate':
                user.is_active = False
                user.save()
                messages.success(request, f'Пользователь {user.username} деактивирован')
    
    return render(request, 'tenders/user_management.html', {'users': users})






@login_required
def chat_room_list(request):
    """Список всех чатов пользователя"""
    # Чаты по тендерам
    tender_rooms = ChatRoom.objects.filter(
        room_type='tender',
        participants=request.user
    ).select_related('tender')
    
    # Чаты по клиентам
    client_rooms = ChatRoom.objects.filter(
        room_type='client',
        participants=request.user
    ).select_related('client')
    
    return render(request, 'tenders/chat_list.html', {
        'tender_rooms': tender_rooms,
        'client_rooms': client_rooms,
    })


@login_required
def chat_room(request, room_id):
    """Комната чата"""
    room = get_object_or_404(ChatRoom, id=room_id, participants=request.user)
    
    # Отметить сообщения как прочитанные
    room.messages.exclude(author=request.user).exclude(read_by=request.user).update()
    for msg in room.messages.exclude(author=request.user).exclude(read_by=request.user):
        msg.read_by.add(request.user)
    
    # Получить сообщения
    messages = room.messages.select_related('author').all()[:100]
    
    # Определить контекст
    context = {
        'room': room,
        'messages': messages,
        'tender': room.tender if room.room_type == 'tender' else None,
        'client': room.client if room.room_type == 'client' else None,
    }
    
    return render(request, 'tenders/chat_room.html', context)


@login_required
@require_POST
def chat_send_message(request, room_id):
    """Отправка сообщения"""
    room = get_object_or_404(ChatRoom, id=room_id, participants=request.user)
    
    content = request.POST.get('content', '').strip()
    
    if content:
        message = ChatMessage.objects.create(
            room=room,
            author=request.user,
            content=content
        )
        
        # Вернуть HTML нового сообщения для AJAX
        message_html = render(request, 'tenders/chat_message.html', {
            'message': message,
            'is_own': True
        }).content.decode('utf-8')
        
        return JsonResponse({
            'success': True,
            'message_html': message_html,
            'message_id': message.id,
        })
    
    return JsonResponse({'success': False, 'error': 'Пустое сообщение'})


@login_required
def chat_get_messages(request, room_id):
    """Получение новых сообщений (для polling)"""
    room = get_object_or_404(ChatRoom, id=room_id, participants=request.user)
    
    last_message_id = request.GET.get('last_id', 0)
    
    messages = room.messages.filter(
        id__gt=last_message_id
    ).select_related('author').order_by('created_at')
    
    messages_html = ''
    for msg in messages:
        is_own = msg.author == request.user
        html = render(request, 'tenders/chat_message.html', {
            'message': msg,
            'is_own': is_own
        }).content.decode('utf-8')
        messages_html += html
    
    return JsonResponse({
        'messages_html': messages_html,
        'last_id': messages.last().id if messages.exists() else last_message_id,
        'unread_count': room.unread_count(request.user),
    })


@login_required
def start_chat(request, content_type, object_id):
    """Создание или открытие чата"""
    if content_type == 'tender':
        tender = get_object_or_404(Tender, id=object_id)
        room, created = ChatRoom.objects.get_or_create(
            room_type='tender',
            tender=tender,
            defaults={}
        )
    elif content_type == 'client':
        client = get_object_or_404(Client, id=object_id)
        room, created = ChatRoom.objects.get_or_create(
            room_type='client',
            client=client,
            defaults={}
        )
    else:
        return redirect('tenders:list')
    
    # Добавить пользователя в участники
    room.participants.add(request.user)
    
    return redirect('tenders:chat_room', room_id=room.id)



    


# ============================================
# 🔹 ФОРМА ЗАЯВКИ НА ОТГРУЗКУ
# ============================================
@login_required
def create_shipment_request(request):
    """Создание заявки на отгрузку"""
    if request.method == 'POST':
        form = ShipmentRequestForm(request.POST)
        formset = ShipmentItemFormSet(request.POST)
        
        if form.is_valid() and formset.is_valid():
            shipment = form.save(commit=False)
            shipment.created_by = request.user
            shipment.save()
            
            # Сохраняем позиции
            formset.instance = shipment
            items = formset.save(commit=False)
            for item in items:
                if item.product:  # Только заполненные позиции
                    item.request = shipment
                    # Автоматический расчёт финальной цены
                    if item.discount_percent:
                        item.final_price = item.unit_price * (1 - item.discount_percent / 100)
                    else:
                        item.final_price = item.unit_price
                    item.save()
            
            # 🔔 Отправляем email менеджеру
            try:
                _send_shipment_notification(shipment)
                messages.success(request, f'✅ Заявка #{shipment.id} создана и отправлена менеджеру!')
            except Exception as e:
                messages.warning(request, f'⚠️ Заявка создана, но email не отправлен: {e}')
            
            return redirect('tenders:shipment_request_detail', pk=shipment.pk)
    else:
        form = ShipmentRequestForm()
        formset = ShipmentItemFormSet()
    
    return render(request, 'tenders/shipment_request_form.html', {
        'form': form,
        'formset': formset,
    })

def _send_shipment_notification(shipment):
    """Отправка email-уведомления о новой заявке"""
    # Получаем email менеджера (или администратора)
    recipient_email = settings.DEFAULT_FROM_EMAIL
    
    # Формируем список позиций
    items_list = []
    for item in shipment.items.all():
        items_list.append(
            f"  • {item.product.name} — {item.quantity} {item.product.unit} "
            f"× {item.final_price} ₽ = {item.total} ₽"
        )
    
    subject = f"🆕 Новая заявка на отгрузку #{shipment.id} от {shipment.client.name}"
    
    message = f"""Здравствуйте!

Поступила новая заявка на отгрузку.

📋 ДЕТАЛИ ЗАЯВКИ:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Номер заявки: #{shipment.id}
Покупатель: {shipment.client.name}
ИНН: {shipment.client.inn or '—'}
Контактное лицо: {shipment.contact_person or '—'}
Телефон: {shipment.contact_phone or '—'}
Email: {shipment.contact_email or '—'}
Дата: {shipment.created_at.strftime('%d.%m.%Y %H:%M')}

📦 ПОЗИЦИИ ЗАЯВКИ:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{chr(10).join(items_list)}

💰 ОБЩАЯ СУММА: {shipment.total_amount:,.2f} ₽
📊 ОБЩЕЕ КОЛИЧЕСТВО: {shipment.total_quantity}

💬 КОММЕНТАРИЙ:
{shipment.comment or '—'}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Просмотреть заявку: http://{settings.ALLOWED_HOSTS[0] if settings.ALLOWED_HOSTS else '127.0.0.1:8000'}/shipment/{shipment.id}/

---
Система SRM Тендеры
"""
    
    send_mail(
        subject=subject,
        message=message,
        from_email=settings.DEFAULT_FROM_EMAIL,
        recipient_list=[recipient_email],
        fail_silently=False,
    )


# ============================================
# 🔹 СПИСОК ЗАЯВОК (РЕЕСТР)
# ============================================
@login_required
def shipment_request_list(request):
    """Реестр заявок на отгрузку"""
    requests_qs = ShipmentRequest.objects.select_related('client', 'created_by').all()
    
    # Фильтрация
    status = request.GET.get('status')
    if status:
        requests_qs = requests_qs.filter(status=status)
    
    client_id = request.GET.get('client')
    if client_id:
        requests_qs = requests_qs.filter(client_id=client_id)
    
    search = request.GET.get('search')
    if search:
        requests_qs = requests_qs.filter(
            Q(client__name__icontains=search) |
            Q(contact_person__icontains=search) |
            Q(id__icontains=search)
        )
    
    return render(request, 'tenders/shipment_request_list.html', {
        'requests': requests_qs,
        'status_choices': ShipmentRequest.STATUS_CHOICES,
        'clients': Client.objects.all(),
    })


# ============================================
# 🔹 ДЕТАЛИ ЗАЯВКИ
# ============================================
@login_required
def shipment_request_detail(request, pk):
    """Детали заявки"""
    shipment = get_object_or_404(
        ShipmentRequest.objects.select_related('client', 'created_by'),
        pk=pk
    )
    items = shipment.items.select_related('product').all()
    
    return render(request, 'tenders/shipment_request_detail.html', {
        'shipment': shipment,
        'items': items,
    })


# ============================================
# 🔹 ИЗМЕНЕНИЕ СТАТУСА ЗАЯВКИ
# ============================================
@login_required
def shipment_request_update_status(request, pk):
    """Изменение статуса заявки"""
    shipment = get_object_or_404(ShipmentRequest, pk=pk)
    
    if request.method == 'POST':
        new_status = request.POST.get('status')
        if new_status in dict(ShipmentRequest.STATUS_CHOICES):
            shipment.status = new_status
            shipment.save()
            messages.success(request, f'Статус заявки #{shipment.id} изменён')
    
    return redirect('tenders:shipment_request_detail', pk=pk)


# ============================================
# 🔹 ЭКСПОРТ РЕЕСТРА В 1С
# ============================================
@login_required
def export_shipment_to_1c(request):
    """Экспорт реестра заявок в Excel для 1С"""
    from django.utils import timezone
    from datetime import datetime
    
    # Получаем ID заявок для экспорта (или все новые)
    request_ids = request.GET.getlist('ids')
    
    if request_ids:
        shipments = ShipmentRequest.objects.filter(id__in=request_ids)
    else:
        # По умолчанию выгружаем все необработанные
        shipments = ShipmentRequest.objects.filter(
            status__in=['new', 'processing', 'approved'],
            exported_to_1c=False
        )
    
    if not shipments.exists():
        messages.warning(request, 'Нет заявок для выгрузки')
        return redirect('tenders:shipment_request_list')
    
    # Создаём Excel файл
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    
    wb = Workbook()
    ws = wb.active
    ws.title = 'Реестр заявок'
    
    # Заголовок
    ws.merge_cells('A1:L1')
    ws['A1'] = f'РЕЕСТР ЗАЯВОК НА ОТГРУЗКУ от {datetime.now().strftime("%d.%m.%Y %H:%M")}'
    ws['A1'].font = Font(bold=True, size=14)
    ws['A1'].alignment = Alignment(horizontal='center')
    
    # Заголовки таблицы
    headers = [
        '№ заявки', 'Дата', 'Покупатель', 'ИНН', 'Контактное лицо',
        'Телефон', 'Email', 'Артикул', 'Наименование продукции',
        'Количество', 'Ед.изм', 'Цена', 'Скидка %', 'Цена с скидкой', 'Сумма', 'Статус'
    ]
    
    header_fill = PatternFill(start_color='0d6efd', end_color='0d6efd', fill_type='solid')
    header_font = Font(bold=True, color='FFFFFF')
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=3, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = thin_border
    
    # Данные
    row = 4
    for shipment in shipments:
        for item in shipment.items.select_related('product'):
            ws.cell(row=row, column=1, value=shipment.id).border = thin_border
            ws.cell(row=row, column=2, value=shipment.created_at.strftime('%d.%m.%Y')).border = thin_border
            ws.cell(row=row, column=3, value=shipment.client.name).border = thin_border
            ws.cell(row=row, column=4, value=shipment.client.inn or '').border = thin_border
            ws.cell(row=row, column=5, value=shipment.contact_person or '').border = thin_border
            ws.cell(row=row, column=6, value=shipment.contact_phone or '').border = thin_border
            ws.cell(row=row, column=7, value=shipment.contact_email or '').border = thin_border
            ws.cell(row=row, column=8, value=item.product.article or '').border = thin_border
            ws.cell(row=row, column=9, value=item.product.name).border = thin_border
            ws.cell(row=row, column=10, value=float(item.quantity)).border = thin_border
            ws.cell(row=row, column=11, value=item.product.unit).border = thin_border
            ws.cell(row=row, column=12, value=float(item.unit_price)).border = thin_border
            ws.cell(row=row, column=13, value=float(item.discount_percent)).border = thin_border
            ws.cell(row=row, column=14, value=float(item.final_price)).border = thin_border
            ws.cell(row=row, column=15, value=float(item.total)).border = thin_border
            ws.cell(row=row, column=16, value=shipment.get_status_display()).border = thin_border
            row += 1
    
    # Автоширина колонок
    for col in ws.columns:
        max_length = 0
        column = col[0].column_letter
        for cell in col:
            try:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            except:
                pass
        ws.column_dimensions[column].width = min(max_length + 2, 30)
    
    # Сохраняем
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    filename = f'registry_1c_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
    response['Content-Disposition'] = f'attachment; filename={filename}'
    wb.save(response)
    
    # Помечаем заявки как выгруженные
    shipments.update(exported_to_1c=True, exported_at=timezone.now())
    
    return response


# ============================================
# 🔹 УПРАВЛЕНИЕ НОМЕНКЛАТУРОЙ
# ============================================
@login_required
def product_list(request):
    """Список продукции"""
    products = Product.objects.filter(is_active=True)
    return render(request, 'tenders/product_list.html', {'products': products})


@login_required
def product_create(request):
    """Добавление продукции"""
    if request.method == 'POST':
        name = request.POST.get('name')
        article = request.POST.get('article', '')
        price = request.POST.get('price', 0)
        unit = request.POST.get('unit', 'шт')
        
        if name:
            Product.objects.create(
                name=name,
                article=article,
                price=price,
                unit=unit
            )
            messages.success(request, f'✅ Продукция "{name}" добавлена')
            return redirect('tenders:product_list')
    
    return render(request, 'tenders/product_form.html')







@login_required
def api_products_prices(request):
    """API: получение цен продукции для автозаполнения"""
    products = Product.objects.filter(is_active=True).values('id', 'name', 'price', 'article', 'unit')
    data = {str(p['id']): p for p in products}
    return JsonResponse(data)